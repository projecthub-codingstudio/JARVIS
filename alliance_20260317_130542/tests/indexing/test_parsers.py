"""Tests for DocumentParser.

Covers text-based formats (existing) and binary document formats
(PDF, DOCX, XLSX, HWPX) per PHASE1_ARCHITECTURE_CORE_DESIGN.md Parser Tiers.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from jarvis.contracts import AccessStatus, DocumentRecord, IndexingStatus
from jarvis.indexing.parsers import DocumentParser


# --- Detect type tests ---


class TestDetectType:
    def test_markdown(self, tmp_path: Path) -> None:
        f = tmp_path / "readme.md"
        f.write_text("# Hello")
        assert DocumentParser().detect_type(f) == "markdown"

    def test_python(self, tmp_path: Path) -> None:
        f = tmp_path / "app.py"
        f.write_text("print('hi')")
        assert DocumentParser().detect_type(f) == "python"

    def test_typescript(self, tmp_path: Path) -> None:
        f = tmp_path / "index.ts"
        f.write_text("const x = 1;")
        assert DocumentParser().detect_type(f) == "typescript"

    def test_yaml(self, tmp_path: Path) -> None:
        f = tmp_path / "config.yaml"
        f.write_text("key: value")
        assert DocumentParser().detect_type(f) == "yaml"

    def test_plain_text(self, tmp_path: Path) -> None:
        f = tmp_path / "notes.txt"
        f.write_text("some notes")
        assert DocumentParser().detect_type(f) == "text"

    def test_unknown_extension(self, tmp_path: Path) -> None:
        f = tmp_path / "data.bin"
        f.write_bytes(b"\x00\x01")
        assert DocumentParser().detect_type(f) == "text"

    def test_pdf(self, tmp_path: Path) -> None:
        f = tmp_path / "report.pdf"
        f.write_bytes(b"%PDF-1.4")
        assert DocumentParser().detect_type(f) == "pdf"

    def test_docx(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.docx"
        f.write_bytes(b"PK")
        assert DocumentParser().detect_type(f) == "docx"

    def test_xlsx(self, tmp_path: Path) -> None:
        f = tmp_path / "sheet.xlsx"
        f.write_bytes(b"PK")
        assert DocumentParser().detect_type(f) == "xlsx"

    def test_hwpx(self, tmp_path: Path) -> None:
        f = tmp_path / "document.hwpx"
        f.write_bytes(b"PK")
        assert DocumentParser().detect_type(f) == "hwpx"


# --- Text parse tests ---


class TestParse:
    def test_parse_markdown(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.md"
        f.write_text("# Title\n\nParagraph text.")
        text = DocumentParser().parse(f)
        assert "Title" in text
        assert "Paragraph text." in text

    def test_parse_python(self, tmp_path: Path) -> None:
        f = tmp_path / "mod.py"
        content = '"""Docstring."""\n\ndef foo():\n    pass\n'
        f.write_text(content)
        text = DocumentParser().parse(f)
        assert "Docstring" in text
        assert "def foo" in text

    def test_parse_nonexistent_raises(self) -> None:
        with pytest.raises(FileNotFoundError):
            DocumentParser().parse(Path("/nonexistent/file.md"))

    def test_parse_empty_file(self, tmp_path: Path) -> None:
        f = tmp_path / "empty.txt"
        f.write_text("")
        assert DocumentParser().parse(f) == ""

    def test_parse_korean_content(self, tmp_path: Path) -> None:
        f = tmp_path / "korean.md"
        f.write_text("# 제목\n\n한국어 문서 내용입니다.")
        text = DocumentParser().parse(f)
        assert "제목" in text
        assert "한국어" in text


# --- PDF parse tests ---


class TestParsePdf:
    def test_parse_pdf_extracts_text(self, tmp_path: Path) -> None:
        """Create a real PDF with PyMuPDF and verify text extraction."""
        import pymupdf

        pdf_path = tmp_path / "test.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello PDF World")
        doc.save(str(pdf_path))
        doc.close()

        text = DocumentParser().parse(pdf_path)
        assert "Hello PDF World" in text

    def test_parse_pdf_korean(self, tmp_path: Path) -> None:
        """PDF with Korean text."""
        import pymupdf

        pdf_path = tmp_path / "korean.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "한국어 PDF 테스트", fontname="helv")
        doc.save(str(pdf_path))
        doc.close()

        text = DocumentParser().parse(pdf_path)
        assert "PDF" in text

    def test_parse_pdf_multipage(self, tmp_path: Path) -> None:
        """Multi-page PDF extracts all pages."""
        import pymupdf

        pdf_path = tmp_path / "multi.pdf"
        doc = pymupdf.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content")
        doc.save(str(pdf_path))
        doc.close()

        text = DocumentParser().parse(pdf_path)
        assert "Page 1 content" in text
        assert "Page 3 content" in text


# --- DOCX parse tests ---


class TestParseDocx:
    def test_parse_docx_extracts_paragraphs(self, tmp_path: Path) -> None:
        """Create a real DOCX and verify paragraph extraction."""
        from docx import Document

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_path))

        text = DocumentParser().parse(docx_path)
        assert "First paragraph" in text
        assert "Second paragraph" in text

    def test_parse_docx_korean(self, tmp_path: Path) -> None:
        """DOCX with Korean text."""
        from docx import Document

        docx_path = tmp_path / "korean.docx"
        doc = Document()
        doc.add_paragraph("한국어 문서입니다")
        doc.add_paragraph("두 번째 단락")
        doc.save(str(docx_path))

        text = DocumentParser().parse(docx_path)
        assert "한국어 문서입니다" in text
        assert "두 번째 단락" in text

    def test_parse_docx_with_table(self, tmp_path: Path) -> None:
        """DOCX table cells are extracted."""
        from docx import Document

        docx_path = tmp_path / "table.docx"
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        doc.save(str(docx_path))

        text = DocumentParser().parse(docx_path)
        assert "Before table" in text
        assert "A1" in text
        assert "B2" in text


# --- XLSX parse tests ---


class TestParseXlsx:
    def test_parse_xlsx_extracts_cells(self, tmp_path: Path) -> None:
        """Create a real XLSX and verify cell extraction."""
        from openpyxl import Workbook

        xlsx_path = tmp_path / "test.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws["A1"] = "Name"
        ws["B1"] = "Value"
        ws["A2"] = "Alpha"
        ws["B2"] = 42
        wb.save(str(xlsx_path))

        text = DocumentParser().parse(xlsx_path)
        assert "[Data]" in text
        assert "Name" in text
        assert "Alpha" in text
        assert "42" in text

    def test_parse_xlsx_korean(self, tmp_path: Path) -> None:
        """XLSX with Korean text."""
        from openpyxl import Workbook

        xlsx_path = tmp_path / "korean.xlsx"
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "이름"
        ws["B1"] = "값"
        ws["A2"] = "테스트"
        wb.save(str(xlsx_path))

        text = DocumentParser().parse(xlsx_path)
        assert "이름" in text
        assert "테스트" in text

    def test_parse_xlsx_multiple_sheets(self, tmp_path: Path) -> None:
        """Multiple sheets are all extracted with sheet names."""
        from openpyxl import Workbook

        xlsx_path = tmp_path / "multi.xlsx"
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1["A1"] = "Data from sheet 1"
        ws2 = wb.create_sheet("Sheet2")
        ws2["A1"] = "Data from sheet 2"
        wb.save(str(xlsx_path))

        text = DocumentParser().parse(xlsx_path)
        assert "[Sheet1]" in text
        assert "[Sheet2]" in text
        assert "Data from sheet 1" in text
        assert "Data from sheet 2" in text


# --- HWPX parse tests ---


class TestParseHwpx:
    def _create_minimal_hwpx(self, path: Path, text_content: str) -> None:
        """Create a minimal HWPX file (ZIP with XML content)."""
        import zipfile

        hwpx_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hs:p>
    <hs:run>
      <hs:t>{text_content}</hs:t>
    </hs:run>
  </hs:p>
</hs:sec>"""

        with zipfile.ZipFile(str(path), "w") as zf:
            zf.writestr("Contents/section0.xml", hwpx_xml)
            zf.writestr("mimetype", "application/hwp+zip")

    def test_parse_hwpx_extracts_text(self, tmp_path: Path) -> None:
        """HWPX text extraction via fallback parser."""
        hwpx_path = tmp_path / "test.hwpx"
        self._create_minimal_hwpx(hwpx_path, "Hello HWPX Document")

        text = DocumentParser().parse(hwpx_path)
        assert "Hello HWPX Document" in text

    def test_parse_hwpx_korean(self, tmp_path: Path) -> None:
        """HWPX with Korean text."""
        hwpx_path = tmp_path / "korean.hwpx"
        self._create_minimal_hwpx(hwpx_path, "한국어 한글 문서입니다")

        text = DocumentParser().parse(hwpx_path)
        assert "한국어" in text
        assert "한글" in text


# --- CreateRecord tests ---


class TestCreateRecord:
    def test_creates_record_with_metadata(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.md"
        f.write_text("content here")
        record = DocumentParser().create_record(f)
        assert isinstance(record, DocumentRecord)
        assert record.path == str(f)
        assert record.size_bytes == f.stat().st_size
        assert record.content_hash  # non-empty SHA-256
        assert record.indexing_status == IndexingStatus.PENDING
        assert record.access_status == AccessStatus.ACCESSIBLE

    def test_record_hash_changes_with_content(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.md"
        f.write_text("version 1")
        r1 = DocumentParser().create_record(f)
        f.write_text("version 2")
        r2 = DocumentParser().create_record(f)
        assert r1.content_hash != r2.content_hash

    def test_record_for_inaccessible_file(self) -> None:
        with pytest.raises(FileNotFoundError):
            DocumentParser().create_record(Path("/nonexistent"))

    def test_record_for_pdf(self, tmp_path: Path) -> None:
        """create_record works for binary files too."""
        import pymupdf

        pdf_path = tmp_path / "test.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "test")
        doc.save(str(pdf_path))
        doc.close()

        record = DocumentParser().create_record(pdf_path)
        assert record.size_bytes > 0
        assert record.content_hash
